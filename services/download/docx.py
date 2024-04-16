import docx as python_docx

from config.logger import log


def extract_text(file_path: str) -> str:
    log().info(f"Extracting text from {file_path}")

    doc = python_docx.Document(file_path)

    paragraphs = []

    log().debug(f"Doc contained: {len(doc.paragraphs)} paragraphs")

    for para_number, para in enumerate(doc.paragraphs, start=1):
        paragraphs.append(para.text)

    return '\n'.join(paragraphs)
